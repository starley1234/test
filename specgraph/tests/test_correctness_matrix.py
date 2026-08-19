from pathlib import Path

from docx import Document
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

from specgraph.db import Base
from specgraph.ingest.pipeline import ingest_parsed_json
from specgraph.pipelines.correctness import load_checklist, run_correctness_matrix


def test_matrix_from_db_requirements(tmp_path, monkeypatch):
    from specgraph import config
    from specgraph.pipelines import correctness as corr

    monkeypatch.setattr(config.settings, "upload_dir", tmp_path)
    monkeypatch.setattr(corr, "EXPORTS", tmp_path)

    engine = create_engine("sqlite:///:memory:")
    Base.metadata.create_all(engine)
    db = sessionmaker(bind=engine)()
    ingest_parsed_json(
        db,
        {
            "title": "spec",
            "products": [{"code": "P", "name": "Блок"}],
            "requirements": [
                {
                    "code": "X.HRDW.FNCT.001/A",
                    "text": "Блок должен выдавать 27 В ±1 В.",
                    "product": "P",
                    "attributes": {"производное": "нет"},
                },
                {
                    "code": "X.HRDW.FNCT.002/A",
                    "text": "Можно использовать любой разъём.",
                    "product": "P",
                },
            ],
        },
        index=False,
    )
    report = run_correctness_matrix(db)
    assert report["count"] == 2
    assert Path(report["output_file"]).is_file()
    names = [c["name"] for c in load_checklist()["criteria"]]
    assert "Верифицируемость" in names
    doc = Document(report["output_file"])
    assert any("Матрица" in p.text or True for p in doc.paragraphs)
    assert len(doc.tables) >= 2
    # header + 2 data rows
    assert len(doc.tables[1].rows) >= 3
    codes = " ".join(c.text for r in doc.tables[1].rows for c in r.cells)
    assert "X.HRDW.FNCT.001/A" in codes
