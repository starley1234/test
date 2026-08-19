"""Пайплайн: методика А.3 → оценка текущих требований → матрица .docx."""

from __future__ import annotations

import json
import re
from datetime import date
from pathlib import Path
from typing import Any
from uuid import uuid4

from sqlalchemy.orm import Session, joinedload

from specgraph.config import settings
from specgraph.llm import invoke_chat
from specgraph.models import Requirement
from specgraph.retrieval.context import expand_requirement

CHECKLIST = Path(__file__).with_name("correctness_checklist.json")
TEMPLATE = Path(__file__).resolve().parents[2] / "input" / "for_pipeline1" / "matrix.docx"
EXPORTS = Path(settings.upload_dir).resolve().parent / "exports"


def load_checklist() -> dict[str, Any]:
    return json.loads(CHECKLIST.read_text(encoding="utf-8"))


def _current_requirements(
    db: Session,
    document_id: int | None,
    product_id: int | None,
    requirement_ids: list[int] | None = None,
) -> list[Requirement]:
    q = (
        db.query(Requirement)
        .options(joinedload(Requirement.attributes))
        .filter(Requirement.is_current.is_(True))
    )
    if document_id:
        q = q.filter(Requirement.document_id == document_id)
    if product_id:
        q = q.filter(Requirement.product_id == product_id)
    if requirement_ids:
        q = q.filter(Requirement.id.in_(requirement_ids))
    rows = [r for r in q.all() if not (r.extra or {}).get("stub") and not (r.extra or {}).get("appendix")]
    return rows


def _heuristic_row(req: Requirement, checklist: dict) -> dict[str, Any]:
    text = (req.text or "").lower()
    attrs = {a.key: a.value for a in req.attributes}
    derived = "да" in (attrs.get("производное") or "").lower()
    has_src = bool(attrs.get("source") or attrs.get("источник требования") or req.parent_id)
    marks = {}
    notes = []
    for c in checklist["criteria"]:
        cid = c["id"]
        mark = "+"
        if cid == "wording" and not re.search(r"должен|необходим|требует|shall", text):
            mark = "–"
            notes.append("нет обязательной формулировки (должен/shall)")
        elif cid == "trace" and not has_src and derived:
            mark = "–"
            notes.append("производное без источника")
        elif cid == "complete" and len(req.text or "") < 20:
            mark = "–"
            notes.append("содержание слишком короткое")
        elif cid == "unamb" and re.search(r"\b(может|рекоменд)", text):
            mark = "–"
            notes.append("необязательная формулировка")
        marks[c["name"]] = mark
    ok = all(v == "+" for v in marks.values())
    return {
        "code": req.code,
        "marks": marks,
        "note": "Несоответствий не обнаружено" if ok else "; ".join(notes) or "есть замечания",
        "pass": ok,
        "mode": "heuristic",
    }


def _llm_row(req: Requirement, checklist: dict, ctx: dict) -> dict[str, Any]:
    crit_txt = "\n".join(
        f"- {c['name']}: " + " ".join(c["questions"][:3]) for c in checklist["criteria"]
    )
    parents = ctx.get("parents") or []
    parent_txt = "\n".join(f"{p.get('code')} stub={p.get('stub')} {(p.get('text') or '')[:200]}" for p in parents)
    attrs = ctx.get("attributes") or {}
    prompt = (
        f"Требование {req.code}\nТекст: {req.text}\n"
        f"Атрибуты: {attrs}\nВышестоящие:\n{parent_txt or 'нет'}\n\n"
        f"Критерии (группа = одна оценка):\n{crit_txt}\n\n"
        "Верни JSON с marks (+/–/н/п) и note. Только JSON."
    )
    raw, usage = invoke_chat(
        "expensive",
        "Ты член группы валидации. Оцени требование по методике рассмотрения корректности. Не выдумывай факты вне карточки.",
        prompt,
    )
    if usage.get("offline") or not raw:
        row = _heuristic_row(req, checklist)
        row["tokens"] = usage
        return row
    m = re.search(r"\{.*\}", raw, re.S)
    data = json.loads(m.group(0) if m else "{}")
    marks = {}
    for c in checklist["criteria"]:
        marks[c["name"]] = str((data.get("marks") or {}).get(c["name"]) or "н/п")[:3]
    note = data.get("note") or "Несоответствий не обнаружено"
    ok = all(v == "+" for v in marks.values())
    return {"code": req.code, "marks": marks, "note": note, "pass": ok, "mode": "llm", "tokens": usage}


def evaluate(
    db: Session,
    *,
    document_id: int | None = None,
    product_id: int | None = None,
    requirement_ids: list[int] | None = None,
    on_progress: Any = None,
) -> dict[str, Any]:
    checklist = load_checklist()
    reqs = _current_requirements(db, document_id, product_id, requirement_ids)
    rows = []
    n = len(reqs)
    if on_progress:
        on_progress({"event": "plan", "step": f"к оценке {n} требований", "total": n})
    for i, r in enumerate(reqs, 1):
        if on_progress:
            on_progress({"event": "item", "step": f"оценка {r.code}", "index": i, "total": n, "code": r.code})
        ctx = expand_requirement(db, r.id)
        row = _llm_row(r, checklist, ctx)
        rows.append(row)
        if on_progress:
            on_progress(
                {
                    "event": "item_done",
                    "step": f"{r.code}: {(row.get('note') or '')[:80]}",
                    "index": i,
                    "total": n,
                    "code": r.code,
                    "tokens": row.get("tokens") or {},
                    "mode": row.get("mode"),
                }
            )
    passed = all(x["pass"] for x in rows) if rows else False
    return {
        "checklist": checklist["title"],
        "result": checklist["pass_text"] if passed else checklist["fail_text"],
        "count": len(rows),
        "rows": rows,
    }


def write_matrix_docx(
    report: dict[str, Any],
    *,
    out_path: Path | None = None,
    designations: dict[str, str] | None = None,
) -> Path:
    from copy import deepcopy

    from docx import Document
    from docx.oxml.ns import qn

    EXPORTS.mkdir(parents=True, exist_ok=True)
    out_path = out_path or EXPORTS / f"matrix_correctness_{date.today().isoformat()}_{uuid4().hex[:6]}.docx"
    src = TEMPLATE if TEMPLATE.exists() else None
    if src:
        doc = Document(str(src))
    else:
        doc = Document()
        doc.add_paragraph("Матрица рассмотрения корректности")
        doc.add_table(rows=5, cols=2)
        doc.add_table(rows=1, cols=15)
        doc.add_table(rows=2, cols=3)

    designations = designations or {}
    if doc.tables:
        meta = doc.tables[0]
        defaults = [
            designations.get("record") or "SpecGraph-auto",
            designations.get("requirements") or "из БД, текущие ревизии",
            designations.get("plan") or "—",
            designations.get("method") or "044.009-0234/02",
            report["result"],
        ]
        for i, val in enumerate(defaults):
            if i < len(meta.rows) and len(meta.rows[i].cells) > 1:
                meta.rows[i].cells[1].text = str(val)

    checklist = load_checklist()
    headers = ["Идентификатор требования/ревизия"] + [c["name"] for c in checklist["criteria"]] + ["Выявленное несоответствие"]
    if len(doc.tables) < 2:
        tbl = doc.add_table(rows=1, cols=len(headers))
    else:
        tbl = doc.tables[1]

    tbl_el = tbl._tbl
    for tr in list(tbl_el.findall(qn("w:tr")))[1:]:
        tbl_el.remove(tr)
    if tbl.rows:
        for i, h in enumerate(headers):
            if i < len(tbl.rows[0].cells):
                tbl.rows[0].cells[i].text = h

    def add_row(values: list[str]) -> None:
        template = tbl.rows[0]._tr
        new_tr = deepcopy(template)
        cells = new_tr.findall(qn("w:tc"))
        for i, cell in enumerate(cells):
            texts = cell.findall(".//" + qn("w:t"))
            val = values[i] if i < len(values) else ""
            if texts:
                texts[0].text = val
                for extra in texts[1:]:
                    extra.text = ""
        tbl_el.append(new_tr)

    for row in report["rows"]:
        vals = [row["code"]]
        for c in checklist["criteria"]:
            vals.append(row["marks"].get(c["name"], "н/п"))
        vals.append(row["note"])
        add_row(vals)

    doc.save(str(out_path))
    return out_path


def run_correctness_matrix(
    db: Session,
    *,
    document_id: int | None = None,
    product_id: int | None = None,
    requirement_ids: list[int] | None = None,
    designations: dict[str, str] | None = None,
    on_progress: Any = None,
    **_: Any,
) -> dict[str, Any]:
    report = evaluate(
        db,
        document_id=document_id,
        product_id=product_id,
        requirement_ids=requirement_ids,
        on_progress=on_progress,
    )
    if on_progress:
        on_progress({"event": "write", "step": "запись matrix.docx"})
    path = write_matrix_docx(report, designations=designations)
    report["output_file"] = str(path)
    report["download"] = f"/exports/{path.name}"
    from specgraph.pipelines.exports import write_text_bundle

    md_lines = [report.get("result") or "", ""]
    for row in report.get("rows") or []:
        md_lines.append(f"- {row.get('code')}: {row.get('note')}")
    extra = write_text_bundle(
        "review-correctness",
        {"output": "\n".join(md_lines), "rows": report.get("rows"), "count": report.get("count")},
        title=report.get("checklist") or "корректность",
    )
    extra["docx"] = report["download"]
    report["downloads"] = extra
    return report
