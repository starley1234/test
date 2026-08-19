"""Выгрузки результатов пайплайна: json / md / xlsx / docx."""

from __future__ import annotations

import json
from datetime import date
from pathlib import Path
from typing import Any
from uuid import uuid4

from specgraph.config import settings

EXPORTS = Path(settings.upload_dir).resolve().parent / "exports"


def write_text_bundle(name: str, payload: dict[str, Any], *, title: str = "") -> dict[str, str]:
    EXPORTS.mkdir(parents=True, exist_ok=True)
    stem = f"{name}_{date.today().isoformat()}_{uuid4().hex[:6]}"
    text = payload.get("output") or payload.get("result") or ""
    if not isinstance(text, str):
        text = json.dumps(text, ensure_ascii=False, indent=2)
    body = {"pipeline": name, "title": title, **{k: v for k, v in payload.items() if k != "downloads"}}

    jp = EXPORTS / f"{stem}.json"
    jp.write_text(json.dumps(body, ensure_ascii=False, indent=2, default=str), encoding="utf-8")
    mp = EXPORTS / f"{stem}.md"
    mp.write_text(f"# {title or name}\n\n{text}\n", encoding="utf-8")

    downloads = {
        "json": f"/exports/{jp.name}",
        "md": f"/exports/{mp.name}",
    }
    try:
        from openpyxl import Workbook

        xp = EXPORTS / f"{stem}.xlsx"
        wb = Workbook()
        ws = wb.active
        ws.title = "result"
        ws.append(["pipeline", name])
        ws.append(["title", title])
        for line in (text or "").splitlines() or [""]:
            ws.append([line])
        wb.save(xp)
        downloads["xlsx"] = f"/exports/{xp.name}"
    except Exception:  # noqa: BLE001
        pass
    try:
        from docx import Document

        dp = EXPORTS / f"{stem}.docx"
        doc = Document()
        doc.add_heading(title or name, 0)
        for para in (text or "").split("\n"):
            doc.add_paragraph(para)
        doc.save(dp)
        downloads["docx"] = f"/exports/{dp.name}"
    except Exception:  # noqa: BLE001
        pass
    return downloads
