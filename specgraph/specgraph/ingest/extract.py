"""Извлечение линейной структуры Word: заголовки, таблицы, картинки.

Каналы:
- python-docx + обход body (стиль Heading N, порядок, drawing/blip);
- JSON скрипта `/extract` (`filename` + `document_structure`);
- Apache Tika для .doc / запасного текста.
"""

from __future__ import annotations

import json
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from specgraph.config import settings


@dataclass
class ExtractedImage:
    filename: str
    content: bytes
    content_type: str
    caption: str | None = None
    rel_id: str | None = None


@dataclass
class Block:
    kind: str  # paragraph | heading | table | image
    text: str = ""
    heading_level: int | None = None
    style: str | None = None
    rows: list[list[str]] = field(default_factory=list)
    entities: list[dict[str, str]] = field(default_factory=list)
    image: ExtractedImage | None = None


@dataclass
class ExtractedDoc:
    title: str | None
    paragraphs: list[str]
    tables: list[list[list[str]]]
    images: list[ExtractedImage] = field(default_factory=list)
    blocks: list[Block] = field(default_factory=list)
    meta: dict[str, Any] = field(default_factory=dict)

    @property
    def text(self) -> str:
        parts = []
        for b in self.blocks:
            if b.kind == "table":
                parts.append("\n".join(" | ".join(r) for r in b.rows))
            elif b.text:
                parts.append(b.text)
        return "\n".join(parts) if parts else "\n".join(self.paragraphs)


def detect_kind(filename: str) -> str:
    lower = filename.lower()
    if lower.endswith(".docm"):
        return "macro_doc"
    if lower.endswith(".docx"):
        return "docx"
    if lower.endswith(".doc"):
        return "doc"
    if lower.endswith(".json"):
        return "parsed_json"
    if lower.endswith((".xlsx", ".xlsm")):
        return "xlsx"
    return "other"


def _ctype(name: str) -> str:
    ext = Path(name).suffix.lower()
    return {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".emf": "image/x-emf",
        ".wmf": "image/x-wmf",
        ".gif": "image/gif",
        ".tif": "image/tiff",
        ".tiff": "image/tiff",
    }.get(ext, "application/octet-stream")


def _heading_level(style_name: str | None) -> int | None:
    if not style_name:
        return None
    s = style_name.strip()
    low = s.lower()
    if low.startswith("heading"):
        tail = low.replace("heading", "").strip()
        if tail.isdigit():
            return int(tail)
    if s.startswith("Заголовок"):
        tail = s.replace("Заголовок", "").strip()
        if tail.isdigit():
            return int(tail)
    return None


def _parse_table_smart(table) -> list[list[str]]:
    table_data: list[list[str]] = []
    for row in table.rows:
        row_cells: list[str] = []
        prev = None
        for cell in row.cells:
            tc = cell._tc
            if tc is prev:
                row_cells.append("")
            else:
                cell_text = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                row_cells.append(cell_text)
            prev = tc
        table_data.append(row_cells)
    return table_data


def _load_media_map(path: Path) -> dict[str, ExtractedImage]:
    """rId → картинка из word/_rels/document.xml.rels + word/media."""
    out: dict[str, ExtractedImage] = {}
    if not zipfile.is_zipfile(path):
        return out
    with zipfile.ZipFile(path) as zf:
        try:
            rels = zf.read("word/_rels/document.xml.rels").decode("utf-8", errors="ignore")
        except KeyError:
            rels = ""
        import re

        for rid, target in re.findall(r'Id="(rId\d+)"[^>]*Target="([^"]+)"', rels):
            if "media/" not in target.replace("\\", "/"):
                continue
            media_path = "word/" + target.replace("../", "")
            if media_path not in zf.namelist():
                # иногда Target уже word/media/...
                alt = target.lstrip("/")
                media_path = alt if alt in zf.namelist() else f"word/{Path(target).name}"
                candidates = [n for n in zf.namelist() if n.endswith(Path(target).name)]
                if candidates:
                    media_path = candidates[0]
            if media_path not in zf.namelist():
                continue
            data = zf.read(media_path)
            fname = Path(media_path).name
            out[rid] = ExtractedImage(filename=fname, content=data, content_type=_ctype(fname), rel_id=rid)
    return out


def _blip_rids(element) -> list[str]:
    from docx.oxml.ns import qn

    rids: list[str] = []
    for blip in element.iter(qn("a:blip")):
        rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
        if rid:
            rids.append(rid)
    return rids


def extract_docx(path: Path) -> ExtractedDoc:
    from docx import Document as Docx
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Docx(str(path))
    media = _load_media_map(path)
    blocks: list[Block] = []
    paragraphs: list[str] = []
    tables: list[list[list[str]]] = []
    images: list[ExtractedImage] = []
    seen_img: set[str] = set()

    body = doc.element.body
    pending_caption: str | None = None

    for element in body:
        tag = element.tag
        if tag == qn("w:p"):
            para = Paragraph(element, doc)
            text = para.text.strip()
            style = para.style.name if para.style is not None else None
            level = _heading_level(style)
            rids = _blip_rids(element)
            if text:
                paragraphs.append(text)
                kind = "heading" if level else "paragraph"
                blocks.append(Block(kind=kind, text=text, heading_level=level, style=style))
                if text.lower().startswith(("рисунок", "рис.", "figure", "схема")):
                    pending_caption = text
            for rid in rids:
                img = media.get(rid)
                if not img or rid in seen_img:
                    continue
                seen_img.add(rid)
                if pending_caption:
                    img.caption = pending_caption
                    pending_caption = None
                images.append(img)
                blocks.append(Block(kind="image", text=img.caption or img.filename, image=img))
        elif tag == qn("w:tbl"):
            tbl = Table(element, doc)
            rows = _parse_table_smart(tbl)
            if rows:
                tables.append(rows)
                blocks.append(Block(kind="table", rows=rows))

    # картинки, не попавшие в body (VML и т.п.)
    for rid, img in media.items():
        if rid not in seen_img:
            images.append(img)
            blocks.append(Block(kind="image", text=img.filename, image=img))

    title = doc.core_properties.title or next((p for p in paragraphs if p), path.stem)
    return ExtractedDoc(
        title=title,
        paragraphs=paragraphs,
        tables=tables,
        images=images,
        blocks=blocks,
        meta={"source": "docx", "images": len(images), "blocks": len(blocks)},
    )


def extract_tika(path: Path) -> ExtractedDoc:
    import httpx

    url = settings.tika_server_url.rstrip("/")
    data = path.read_bytes()
    try:
        r = httpx.put(f"{url}/tika", content=data, headers={"Accept": "text/plain"}, timeout=60)
        r.raise_for_status()
        text = r.text
        meta_r = httpx.put(f"{url}/meta", content=data, headers={"Accept": "application/json"}, timeout=60)
        meta = meta_r.json() if meta_r.status_code == 200 else {}
    except Exception as exc:  # noqa: BLE001
        return ExtractedDoc(title=path.stem, paragraphs=[], tables=[], meta={"tika_error": str(exc)})

    paragraphs = [ln.strip() for ln in text.splitlines() if ln.strip()]
    blocks = [Block(kind="paragraph", text=p) for p in paragraphs]
    return ExtractedDoc(
        title=meta.get("title") or path.stem,
        paragraphs=paragraphs,
        tables=[],
        blocks=blocks,
        meta=meta,
    )


def extract_parsed_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def extracted_from_script_json(payload: dict[str, Any]) -> ExtractedDoc:
    """JSON ответа скрипта pars.txt / POST /extract."""
    structure = payload.get("document_structure") or payload.get("structure") or []
    blocks: list[Block] = []
    paragraphs: list[str] = []
    tables: list[list[list[str]]] = []
    for item in structure:
        if not isinstance(item, dict):
            continue
        typ = (item.get("type") or "paragraph").lower()
        content = item.get("content")
        ents = item.get("entities") or []
        if typ == "table" and isinstance(content, list):
            rows = [[str(c) if c is not None else "" for c in row] for row in content if isinstance(row, list)]
            tables.append(rows)
            blocks.append(Block(kind="table", rows=rows, entities=ents))
        else:
            text = content if isinstance(content, str) else str(content or "")
            text = text.strip()
            if not text:
                continue
            paragraphs.append(text)
            level = _heading_level(None)
            # эвристика заголовка: коротко, без точки, или совпало с известными разделами
            kind = "paragraph"
            hl = None
            if len(text) < 90 and not text.endswith(".") and "\n" not in text:
                kind = "heading"
                hl = 1
            blocks.append(Block(kind=kind, text=text, heading_level=hl, entities=ents))
    title = payload.get("title") or payload.get("filename") or "document"
    return ExtractedDoc(
        title=title,
        paragraphs=paragraphs,
        tables=tables,
        blocks=blocks,
        meta={"source": "script_json", "filename": payload.get("filename")},
    )


def extract_xlsx(path: Path) -> ExtractedDoc:
    """Таблица Excel как набор table-блоков (приложение к требованию)."""
    try:
        import openpyxl
    except ImportError:
        return ExtractedDoc(title=path.stem, paragraphs=[path.stem], tables=[], meta={"xlsx_error": "openpyxl missing"})
    wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
    blocks: list[Block] = []
    tables: list[list[list[str]]] = []
    paragraphs = [path.stem]
    blocks.append(Block(kind="heading", text=path.stem, heading_level=1))
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        blocks.append(Block(kind="heading", text=f"Лист {sheet}", heading_level=2))
        rows: list[list[str]] = []
        for row in ws.iter_rows(values_only=True):
            cells = ["" if c is None else str(c).strip() for c in row]
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append(rows)
            blocks.append(Block(kind="table", rows=rows))
    wb.close()
    return ExtractedDoc(
        title=path.stem,
        paragraphs=paragraphs,
        tables=tables,
        blocks=blocks,
        meta={"source": "xlsx", "sheets": list(wb.sheetnames) if False else []},
    )


def extract_any(path: Path) -> ExtractedDoc:
    kind = detect_kind(path.name)
    if kind == "xlsx":
        return extract_xlsx(path)
    if kind == "parsed_json":
        payload = extract_parsed_json(path)
        if "document_structure" in payload:
            return extracted_from_script_json(payload)
        # произвольный JSON — как текст секций
        from specgraph.ingest.structure import extracted_from_generic_json

        return extracted_from_generic_json(payload)
    if kind in {"docx", "macro_doc"}:
        extracted = extract_docx(path)
        if kind == "macro_doc":
            extra = extract_tika(path)
            if extra.paragraphs and len(extra.paragraphs) > len(extracted.paragraphs):
                extracted.meta["tika_extra"] = extra.meta
        return extracted
    if kind == "doc":
        return extract_tika(path)
    tika = extract_tika(path)
    if tika.paragraphs:
        return tika
    text = path.read_text(encoding="utf-8", errors="ignore")
    paras = [ln.strip() for ln in text.splitlines() if ln.strip()]
    return ExtractedDoc(
        title=path.stem,
        paragraphs=paras,
        tables=[],
        blocks=[Block(kind="paragraph", text=p) for p in paras],
    )
